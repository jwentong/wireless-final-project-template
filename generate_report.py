"""Generate the final project report (docx) following the teacher's template.

Builds cover info table, academic-integrity statement, scoring table, abstract,
keywords and the main body (overview / design / mock / implementation /
experiments / AI workflow / conclusion) with the six result figures embedded.
Run: python generate_report.py  ->  报告.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

RESULTS = Path("results")
OUTPUT = Path("报告.docx")

# Identity placeholders — fill before submission (GitHub username known from fork)
STUDENT_ID = "2024140044"
STUDENT_NAME = "齐宇恒"
GITHUB_USER = "EchoTreee"
FORK_URL = "https://github.com/EchoTreee/wireless-final-project-template"


def _set_base_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")


def h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def p(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)


def figure(doc: Document, name: str, caption: str, width: float = 5.6) -> None:
    path = RESULTS / name
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v


def build() -> None:
    doc = Document()
    _set_base_font(doc)

    # ---------------- Cover ----------------
    title = doc.add_heading("深圳大学 无线通信技术期末项目报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("基于 AI 辅助编程的无线通信文件传输基带仿真系统")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    kv_table(doc, [
        ("课程名称", "无线通信技术"),
        ("项目题目", "无线通信文件传输基带仿真系统"),
        ("学号", STUDENT_ID),
        ("姓名", STUDENT_NAME),
        ("GitHub 用户名", GITHUB_USER),
        ("学生 Fork 仓库地址", FORK_URL),
        ("提交分支", "main"),
    ])

    # ---------------- Integrity statement ----------------
    doc.add_page_break()
    h(doc, "学术诚信承诺书", 1)
    p(doc, "本人在此声明，所提交的课程期末项目报告《基于 AI 辅助编程的无线通信文件传输基带仿真系统》"
           "是本人在遵守课程要求和学术规范前提下独立完成的，具有原创性，未抄袭、剽窃他人成果或侵犯他人知识产权。")
    for item in [
        "1. 报告中的系统设计、代码说明、测试分析、结果解释和观点均源自本人对 PRD、课程知识和实验结果的理解与分析。",
        "2. 对他人资料、开源代码、文献、AI 工具输出等引用，已在报告或 AI_LOG.md 中说明并标明来源。",
        "3. 报告中使用的文献、资料、代码片段等来源均已列出并尽力准确说明。",
        "4. 本人明确知晓学术不端行为的严重性，承诺不抄袭、不伪造结果、不硬编码公开测试、不绕过通信链路直接复制文件。",
        "5. 本人理解并接受课程关于 AI 辅助编程的要求，保留 AI_LOG.md，并能解释每个模块的通信原理、参数、逻辑与结果。",
    ]:
        p(doc, item)
    sign = doc.add_paragraph("学生签名：____________    日期：______年____月____日")
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ---------------- Scoring table ----------------
    doc.add_page_break()
    h(doc, "期末项目评分表", 1)
    scoring = [
        ("评分项", "分值", "得分"),
        ("需求理解与设计文档", "20", ""),
        ("mock 测试与设计修订", "15", ""),
        ("系统代码实现", "25", ""),
        ("公开与隐藏测试", "20", ""),
        ("实验结果与分析", "10", ""),
        ("AI 使用记录与工程规范", "10", ""),
        ("总分", "100", ""),
    ]
    table = doc.add_table(rows=len(scoring), cols=3)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(scoring):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val

    # ---------------- Abstract ----------------
    doc.add_page_break()
    h(doc, "摘要", 1)
    p(doc, "本项目基于 AI 辅助编程实现了一个端到端的无线通信文件传输基带仿真系统，"
           "将 UTF-8 文本 Test.txt 经源编码、PN 序列扰码、卷积码信道编码、帧封装、QPSK 调制、"
           "AWGN 信道、前导相关同步、解调、Viterbi 译码、解扰与源解码后恢复为 received.txt。"
           "系统严格遵循 PRD 的统一口径（Gray 映射 QPSK、符号功率 SNR 定义、length 字段语义、"
           "0–128 符号同步偏移），并采用 PRD→DESIGN→TEST_PLAN→mock→TDD→验证 的规范工程流程。"
           "mock 测试发现帧头长度字段不受前向纠错保护的设计缺陷并据此修订（3× 重复保护）。"
           "在 SNR≥12 dB、AWGN、固定随机种子下，received.txt 与 Test.txt 完全一致，"
           "误比特率为 0，文本恢复率为 1.0，校验通过。系统进一步实现卷积码 Viterbi 译码、"
           "Rayleigh/Rician 衰落与 ZF/MMSE 均衡、BPSK/QPSK/16-QAM 多调制对比等提高模块，"
           "并通过 BER-SNR 曲线、星座图、同步相关峰图量化分析了编码增益、均衡价值与系统瓶颈。"
           "单元测试 47 项、公开测试 22 项全部通过。")
    h(doc, "关键词", 2)
    p(doc, "无线通信；QPSK；AWGN；帧同步；卷积码；Viterbi 译码；信道均衡；AI 辅助编程")

    # ---------------- 1 Overview ----------------
    doc.add_page_break()
    h(doc, "1 项目概述", 1)
    h(doc, "1.1 项目背景与目标", 2)
    p(doc, "无线通信系统需要在存在噪声、衰落与同步偏差的信道上可靠传输信息。本项目以课程 PRD 为需求基线，"
           "实现一个可运行的无线通信基带仿真系统，将教师给定的约 300 字中文课程描述 Test.txt 作为业务载荷，"
           "经过发送端、无线信道与接收端处理后在 received.txt 恢复原文，并输出误比特率、帧错误率、"
           "文本恢复率等性能指标与可视化结果，重点考查对通信链路的整体理解与模块化实现能力。")
    h(doc, "1.2 PRD 关键需求理解", 2)
    p(doc, "系统实现 PRD 规定的固定链路（顺序不可变）：Source Encode → Encrypt/Scramble → Channel Encode → "
           "Frame Build → QPSK Modulate → Channel → Synchronization → QPSK Demodulate → Channel Decode → "
           "Decrypt/Descramble → Source Decode → received.txt → Metrics/Plots。统一命令行入口为 "
           "python main.py --input Test.txt --output results/received.txt --snr 12 --seed 2026 --mod qpsk --channel awgn，"
           "运行后生成 results/received.txt、results/metrics.json 及至少两张图表。验收基线为 SNR≥12 dB、AWGN、"
           "固定 seed 下 received.txt 与 Test.txt 完全一致。")
    h(doc, "1.3 GitHub 提交信息", 2)
    kv_table(doc, [
        ("GitHub 用户名", GITHUB_USER),
        ("Fork 仓库地址", FORK_URL),
        ("提交分支", "main"),
        ("统一运行命令", "python main.py --input Test.txt --output results/received.txt --snr 12 --seed 2026 --mod qpsk --channel awgn"),
    ])

    # ---------------- 2 Design ----------------
    h(doc, "2 系统设计", 1)
    h(doc, "2.1 总体架构与固定链路", 2)
    p(doc, "系统按模块化原则拆分为 src/ 下的独立模块：source（源编码）、scramble（扰码）、channel_coding（信道编码）、"
           "framing（帧）、modulation（调制）、channel（信道）、synchronization（同步）、equalizer（均衡）、"
           "metrics（指标）、pipeline（端到端编排）。各模块函数命名与公开测试的自动发现约定一致，保证零适配通过。")
    h(doc, "2.2 帧结构设计", 2)
    p(doc, "帧格式（比特域，MSB-first）：前导(Barker-13，26 bit) | orig_len(32 bit×3 重复保护) | "
           "coded_len(32 bit×3 重复保护) | payload | CRC-16 | padding。其中 orig_len 表示源编码后、扰码前的"
           "原始 payload 比特数，用于接收端去 padding 恢复 UTF-8；coded_len 表示信道编码后进帧的 payload 比特数，"
           "用于定位 payload 与校验字段。将两者拆分为独立字段，解决了 PRD 中 length 语义与接收端定位需求的冲突。"
           "前导取 Barker-13 序列（+1→比特 00、-1→比特 11），经 QPSK 调制后落在主对角线，自相关旁瓣≤1，同步峰尖锐。")
    h(doc, "2.3 QPSK 调制与归一化", 2)
    p(doc, "基础调制采用 PRD 强制的 Gray 映射 QPSK：00→(1+j)/√2、01→(-1+j)/√2、11→(-1-j)/√2、10→(1-j)/√2。"
           "Gray 编码使相邻象限仅差 1 比特，降低误符号到误比特的传播；归一化因子 1/√2 使符号平均功率为 1。"
           "QPSK 可视为两路正交的 BPSK，解调时按 I、Q 两路实部/虚部符号位独立判决。")
    h(doc, "2.4 SNR 定义与 Eb/N0 换算", 2)
    p(doc, "SNR 定义为接收端符号平均功率与复高斯噪声平均功率之比（dB），与 PRD 一致。给定 snr_db，"
           "噪声平均功率 Pn = Ps/10^(snr_db/10)，实部、虚部各 Pn/2。QPSK 每符号承载 2 比特，故 "
           "Eb/N0 = SNR/log2(4) = SNR/2，即 Eb/N0(dB) = SNR(dB) − 3.01 dB。该换算用于与理论 BER 曲线对照。")
    h(doc, "2.5 信道编码：卷积码 K=7 + Viterbi", 2)
    p(doc, "默认信道编码采用约束长度 K=7、码率 1/2、生成多项式 (171,133)₈ 的卷积码，发送端附加 6 个零尾比特"
           "使编码器回到全零态，接收端用硬判决 Viterbi 算法在 64 状态网格上做加比选并回溯最优路径。"
           "相比无编码，卷积码可在相同 BER 下降低所需 Eb/N0 约 4–5 dB（编码增益），是低 SNR 鲁棒性的主要来源。"
           "系统同时实现 Hamming(7,4) 作为对比基线。")

    # ---------------- 3 Mock ----------------
    h(doc, "3 mock 测试与设计修订", 1)
    p(doc, "在正式实现前，先用最小桩实现（信道编码以重复码占位）搭建端到端 mock 骨架，验证接口、帧结构、"
           "同步与端到端流程是否可行，并通过压低 SNR 的鲁棒性扫描尽早暴露设计缺陷。")
    h(doc, "3.1 发现的设计缺陷", 2)
    p(doc, "mock 鲁棒性扫描显示：低 SNR 下系统大面积失败，且失败集中于"
           "“帧头长度字段损坏”。根因是 PRD 固定链路中信道编码位于帧封装之前，导致帧头的 length 字段在"
           "前向纠错保护范围之外——一个比特错误落在 32 位长度字段上即导致整帧解析错位。修订前在 4 dB 有 18/20、"
           "6 dB 有 11/20 属此类失败。")
    h(doc, "3.2 设计修订与验证", 2)
    p(doc, "修订方案：对 orig_len/coded_len 长度字段采用 3× 重复 + 多数判决保护（不改变 PRD 链路顺序，"
           "每帧仅增加 128 比特开销）。验证脚本对比表明，header 字段损坏率在 6 dB 由 11→1、4 dB 由 18→6 显著下降，"
           "失败模式转移为可由正式卷积码进一步纠正的 payload 残余误码。该修订已回填 DESIGN.md（v0.2）。")

    # ---------------- 4 Implementation ----------------
    h(doc, "4 系统实现", 1)
    h(doc, "4.1 关键算法", 2)
    p(doc, "（1）卷积码/Viterbi：预计算网格状态转移与输出，硬判决汉明度量加比选，零尾终止保证无噪声完全可逆。"
           "（2）CRC 分层：校验 CRC 覆盖原始信息位并随信息一起进入信道编码受 FEC 保护，接收端译码后再校验，"
           "符合 PRD‘校验覆盖原始 payload bitstream’的要求且纠错后能通过。"
           "（3）同步：对任意前导做归一化滑动互相关取峰，覆盖 0–128 符号偏移。"
           "（4）均衡：衰落信道下用前导估计信道增益 h，再以 MMSE 均衡。")
    h(doc, "4.2 实现期关键调试", 2)
    p(doc, "调试一：首版 CRC 算在信道编码后的 coded payload 上，AWGN 12 dB 下该比特序列存在被 Viterbi 纠正前的"
           "残余错误，导致校验失败而文本却完全恢复，二者矛盾；修复为 CRC 覆盖原始信息位并受 FEC 保护。"
           "调试二：均衡单元测试中噪声功率定义与全系统 SNR 约定不一致，导致 MMSE 过度收缩；统一定义后 MMSE 恢复理论最优。")

    # ---------------- 5 Experiments ----------------
    h(doc, "5 测试与实验结果分析", 1)
    p(doc, "测试体系分为单元测试（47 项）、集成与端到端测试、以及教师公开测试（22 项），全部通过。"
           "在 SNR≥12 dB、AWGN、seed=2026 下，received.txt 与 Test.txt 完全一致，metrics 记录 "
           "ber=0、text_match_rate=1.0、checksum_pass=true。")
    h(doc, "5.1 QPSK 星座图", 2)
    figure(doc, "constellation.png", "图 1 SNR=12 dB AWGN 下接收 QPSK 星座图，四象限四簇清晰可分。")
    p(doc, "星座图直观解释了误码来源：高 SNR 下四个象限的点云紧凑、远离判决边界；当 SNR 降低、噪声云扩散并越过"
           "I/Q 轴判决边界时即产生误比特。")
    h(doc, "5.2 同步相关峰", 2)
    figure(doc, "sync_peak.png", "图 2 前导归一化互相关曲线，峰值位置即检测到的帧起点。")
    p(doc, "前导的尖锐自相关使相关曲线在帧起点处形成明显峰值，旁瓣低，在 SNR≥12 dB 下帧起点检测误差不超过 1 个符号。")
    h(doc, "5.3 BER-SNR 与编码增益", 2)
    figure(doc, "ber_curve.png", "图 3 QPSK 无编码与卷积码 BER-SNR 曲线。")
    figure(doc, "fec_comparison.png", "图 4 无编码 / Hamming(7,4) / 卷积码 BER 对比。")
    p(doc, "随 SNR 升高 BER 单调下降；卷积码曲线显著陡于无编码与汉明码，体现约 4–5 dB 的编码增益。"
           "在 10 dB 处，QPSK 无编码 BER≈5×10⁻⁴，卷积码已降至 0（实验样本内无误码）。")
    h(doc, "5.4 多调制对比", 2)
    figure(doc, "modulation_comparison.png", "图 5 BPSK / QPSK / 16-QAM 在 AWGN 下的 BER 对比。")
    p(doc, "在相同符号 SNR 口径下，调制阶数越高 BER 越差但频谱效率越高：BPSK 最优、16-QAM 最差。"
           "其中 QPSK 较 BPSK 约差 3 dB，是因为 QPSK 每符号承载 2 比特，相同符号 SNR 下 Eb/N0 低 3 dB；"
           "若按 Eb/N0 口径，BPSK 与 QPSK 的 BER 理论上相同。这一现象自适应调制的依据：高 SNR 选高阶调制提速、"
           "低 SNR 退回低阶调制保可靠。")
    h(doc, "5.5 衰落信道与均衡", 2)
    figure(doc, "fading_comparison.png", "图 6 AWGN 与 Rayleigh 衰落（无均衡 / ZF / MMSE）BER 对比。")
    p(doc, "Rayleigh 衰落在不均衡时 BER 居高不下（信道随机旋转/缩放使判决失效）；经均衡后 BER 随 SNR 下降。"
           "值得注意的是，对仅作相位判决的 QPSK，ZF 与 MMSE 仅相差一个正实标量，硬判决 BER 完全重合（图中 ZF 以虚线叠示）；"
           "而在估计均方误差意义下 MMSE 优于 ZF，这一点已由单元测试验证。两个角度并不矛盾，分别对应硬判决与软估计两种度量。")
    h(doc, "5.6 失败与误码原因分析", 2)
    p(doc, "当 SNR 降低时，系统中最先失效的环节依次为：同步（相关峰被噪声淹没导致帧起点错位、整帧乱码）与"
           "信道译码（噪声超出卷积码纠错能力产生残余误码）。若 received.txt 出现乱码，排查顺序为："
           "先看同步相关峰是否定位正确 → 校验 checksum_pass → 观察 BER → 检查星座扩散程度，逐级定位瓶颈。")

    # ---------------- 6 AI workflow ----------------
    h(doc, "6 AI 辅助编程与工程流程", 1)
    p(doc, "项目采用 Claude Code + Superpowers，严格按 Vibe→Spec→Harness 思想推进："
           "先以最低成本澄清需求（PRD 理解与满分策略），再将隐含判断收敛为显式工程约束（DESIGN.md、TEST_PLAN.md），"
           "随后以 mock 验证设计、以 TDD 实现并以子代理/测试做完成前验证。完整的关键 prompt、AI 生成内容、"
           "人工修改与采纳理由记录于 AI_LOG.md。AI 生成内容中，整体架构与接口被保留；length 位宽、帧头保护、"
           "CRC 分层、均衡 SNR 定义等被人工修改；‘直接生成最终代码’的做法被拒绝，以保证流程留痕与可解释性。"
           "全部代码均可由本人解释通信原理、参数与逻辑，无针对公开测试的硬编码，无绕过链路的文件直拷。")

    # ---------------- 7 Conclusion ----------------
    h(doc, "7 结论", 1)
    p(doc, "本项目实现了符合 PRD 统一口径的端到端无线通信基带仿真系统，在 SNR≥12 dB、AWGN、固定 seed 下"
           "实现文本无损恢复，单元测试与公开测试全部通过，并通过卷积码 Viterbi、Rayleigh/Rician 均衡、"
           "多调制对比等提高模块达到 Level 3 要求。规范的工程流程（含 mock 驱动的设计修订）与完整的 AI_LOG"
           "使系统的每一处设计都可追溯、可解释。")

    h(doc, "参考文献", 1)
    for ref in [
        "[1] J. G. Proakis, M. Salehi. Digital Communications. McGraw-Hill.",
        "[2] A. J. Viterbi. Error bounds for convolutional codes and an asymptotically optimum decoding algorithm. IEEE Trans. IT, 1967.",
        "[3] 无线通信技术期末项目 PRD 与公开测试集（教师提供）。",
        "[4] Claude Code Technical Documents. https://code.claude.com/docs",
    ]:
        p(doc, ref)

    doc.save(OUTPUT)
    print(f"saved: {OUTPUT.resolve()}  (paragraphs={len(doc.paragraphs)})")


if __name__ == "__main__":
    build()
